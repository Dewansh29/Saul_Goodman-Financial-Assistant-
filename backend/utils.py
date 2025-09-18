import fitz
import camelot
import pandas as pd
from typing import List, Dict, Any
import io
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- FIX: The Excel processing logic is now much more robust ---
def process_excel_data(dataframes: Dict[str, pd.DataFrame]) -> str:
    """
    Scans through Excel sheets, finds rows with financial keywords across ALL columns,
    and returns a clean string of relevant data for the LLM.
    """
    relevant_data_text = ""
    keywords = ['revenue', 'income', 'expenses', 'profit', 'ebitda', 'assets', 'liabilities', 'equity', 'cash']

    for sheet_name, df in dataframes.items():
        df.columns = [str(col) if 'unnamed' not in str(col).lower() else '' for col in df.columns]
        df_str = df.astype(str).apply(lambda x: x.str.lower())

        relevant_rows = []
        for index, row in df_str.iterrows():
            # CHANGE: Search the ENTIRE row for keywords, not just the first two columns.
            if len(row) > 0 and any(keyword in str(cell) for keyword in keywords for cell in row):
                original_row = df.iloc[index]
                row_text = " | ".join([f"{col}: {val}" for col, val in original_row.items() if pd.notna(val) and str(val).strip() != ''])
                if row_text:
                    relevant_rows.append(row_text)

        if relevant_rows:
            relevant_data_text += f"--- Relevant Data from Sheet: {sheet_name} ---\n"
            relevant_data_text += "\n".join(relevant_rows)
            relevant_data_text += "\n\n"
            
    return relevant_data_text if relevant_data_text else "No financial keywords found in the Excel file."

def extract_tables_from_pdf(pdf_bytes: bytes, page_numbers: List[int]) -> Dict[int, List[pd.DataFrame]]:
    tables_by_page = {}
    if not page_numbers:
        return tables_by_page
    pdf_file = io.BytesIO(pdf_bytes)
    pages_str = ",".join(map(str, page_numbers))
    print(f"Attempting to extract tables from pages using Camelot: {pages_str}")
    try:
        tables = camelot.read_pdf(pdf_file, pages=pages_str, flavor='stream', line_scale=40)
        for table_report in tables:
            page = table_report.page
            if page not in tables_by_page:
                tables_by_page[page] = []
            tables_by_page[page].append(table_report.df)
            print(f"Found a table on page {page} with Camelot.")
        return tables_by_page
    except Exception as e:
        print(f"An error occurred while processing with Camelot: {e}")
        return {}

def extract_text_from_pdf(pdf_bytes: bytes, page_numbers: List[int]) -> str:
    full_text = ""
    pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page_num in page_numbers:
        if 1 <= page_num <= len(pdf_document):
            page = pdf_document.load_page(page_num - 1)
            full_text += page.get_text("text") + "\n---\n"
    return full_text

def extract_data_from_excel(file_bytes: bytes, filename: str) -> Dict[str, pd.DataFrame]:
    print(f"--- Attempting to extract data from spreadsheet: {filename} ---")
    dataframes = {}
    try:
        if filename.endswith('.csv'):
            df = pd.read_csv(io.BytesIO(file_bytes))
            dataframes['sheet_1'] = df
        elif filename.endswith(('.xls', '.xlsx')):
            xls = pd.ExcelFile(io.BytesIO(file_bytes), engine='openpyxl')
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                dataframes[sheet_name] = df
        print(f"--- Successfully extracted {len(dataframes)} sheet(s) ---")
        return dataframes
    except Exception as e:
        print(f"An error occurred while processing the spreadsheet: {e}")
        return {}

def create_word_report(report_data: Dict, company_name: str) -> io.BytesIO:
    try:
        document = Document('template.docx')
    except Exception:
        print("WARNING: template.docx not found. Creating report without watermark.")
        document = Document()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Saul Goodman: Strategy & Analysis')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    tagline = document.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run('Better Call Saul')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.italic = True
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    
    document.add_paragraph('---' * 25)
    
    document.add_heading(f'Subject: Financial Analysis of {company_name}', level=2)
    
    document.add_heading('1. Executive Summary & Final Verdict', level=1)
    summary_text = report_data.get('final_summary', 'No summary was generated.')
    document.add_paragraph(summary_text)
    
    document.add_heading('2. Key Financial Metrics Extracted', level=1)
    kpi_data_str = report_data.get('cleaned_data', '{}')
    if kpi_data_str and kpi_data_str != '{}' and "Error" not in kpi_data_str:
        try:
            kpi_data = json.loads(kpi_data_str)
            for key, value in kpi_data.items():
                p = document.add_paragraph(style='List Bullet')
                p.add_run(f'{key}: ').bold = True
                p.add_run(str(value) if value is not None else 'N/A')
        except:
            document.add_paragraph("Could not format the extracted KPI data.")
    else:
        document.add_paragraph("No structured financial KPIs were extracted from the document.")
            
    document.add_heading('3. The Boardroom: Analyst Debate Transcript', level=1)
    debate_entries = report_data.get('debate', [])
    if not debate_entries:
        document.add_paragraph("No debate was generated.")
    else:
        for entry in debate_entries:
            p = document.add_paragraph(style='List Bullet')
            if ':' in entry:
                persona, text = entry.split(':', 1)
                p.add_run(f'{persona}:').bold = True
                p.add_run(text)
            else:
                p.add_run(entry)
        
    document.add_paragraph("\n" + "---" * 25)
    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("CONFIDENTIAL MEMORANDUM | Generated by the Saul Goodman Financial Strategy Assistant.")
    run.font.size = Pt(10)
    run.italic = True
    
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    
    return file_stream